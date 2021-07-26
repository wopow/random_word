#  -*-coding:utf8 -*-
import random
import os
from natsort import natsorted
from docx import Document
from docx.oxml.ns import qn

def make_file(file):
    outfile_name=file.split('.')[0]+'.docx'
    outfile1_name=file.split('.')[0]+'答案.docx'
#    outfile = open(outfile_name, 'w')
#    outfile1 = open(outfile1_name, 'w')
    outfile = Document()
    outfile.styles['Normal'].font.name = u'宋体'
    outfile.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    section = outfile.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2')
    outfile1 = Document()
    outfile1.styles['Normal'].font.name = u'宋体'
    outfile1.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    section1 = outfile1.sections[0]
    sectPr1 = section1._sectPr
    cols1 = sectPr1.xpath('./w:cols')[0]
    cols1.set(qn('w:num'), '2')
    file = open('english\\'+file, 'r', encoding='utf-8')
    words = file.readlines()
    word_dict = {}
    for i in words:
        word_dict[i.split(':')[0]] = i.split(':')[1].replace('\n', '')
    word_dict_key = list(word_dict.keys())
    random.shuffle(word_dict_key)
    new_word_dict = {}
    for key in word_dict_key:
        new_word_dict[key] = word_dict.get(key)
        new_line = word_dict.get(key).ljust(len(key) + len(word_dict.get(key)) + 10, '_')
        outfile.add_paragraph(word_dict.get(key) + ' :' + key)
        outfile1.add_paragraph(new_line)
    outfile1.save(outfile_name)
    outfile.save(outfile1_name)
    file.close()

rootdir = 'english'
dirlist = os.listdir(rootdir)
dirlist = natsorted(dirlist)
number = input('请输入要打印的单元：')
start = int(number.split('-')[0])
end = int(number.split('-')[1])
for i in range(start - 1, end):
    make_file(dirlist[i])
